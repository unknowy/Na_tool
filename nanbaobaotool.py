#-*- coding:utf-8 -*-
import os
import xlrd
import zipfile
import time
import sys
import shutil
from xlrd import open_workbook
from xlutils.copy import copy
from docx import Document
from docx.shared import Pt
from docx.shared import Inches
from docx.oxml.ns import qn


def doc(word):
    document = Document(u'report.docx')
    paragraph = document.add_paragraph('\n')
    run = paragraph.add_run(word)
    run.font.name = u'宋体'
    r = run._element
    r.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
    document.save(u'report.docx')


def doc_table_risk_data1(excel_name):
    book = xlrd.open_workbook(excel_name)
    sheet_name1 = book.sheet_names()[1]  # 获得指定索引的sheet名字
    sheet = book.sheet_by_name(sheet_name1)
    row = 2
    date = sheet.row_values(row)
    while ((date[5] + date[6]) != 0):
        date = sheet.row_values(row + 1)
        row = row + 1

    document = Document('report.docx')
    document.add_page_break()
    table = document.add_table(rows=row - 1, cols=10)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = u'序号'
    hdr_cells[1].text = u'IP地址'
    hdr_cells[2].text = u'所属部门'
    hdr_cells[3].text = u'使用人'
    hdr_cells[4].text = u'操作系统'
    hdr_cells[5].text = u'风险等级'
    hdr_cells[6].text = u'高危漏洞'
    hdr_cells[7].text = u'中危漏洞'
    hdr_cells[8].text = u'合计'
    hdr_cells[9].text = u'主机风险值'

    for i in xrange(1, row):

        date = sheet.row_values(i + 1)
        if((date[5] + date[6]) == 0):
            break
        row_cells = table.rows[i].cells
        row_cells[0].text = str(i)
        row_cells[1].text = str(date[1])
        row_cells[4].text = date[3]
        row_cells[5].text = date[4]
        row_cells[6].text = str(date[5])
        row_cells[7].text = str(date[6])
        row_cells[8].text = str(date[5] + date[6])
        row_cells[9].text = str(date[9])

    document.save(u'report.docx')


def doc_table_risk_data2(excel_name):
    book = xlrd.open_workbook(excel_name)
    sheet_name1 = book.sheet_names()[2]  # 获得指定索引的sheet名字
    sheet = book.sheet_by_name(sheet_name1)
    row = 2
    date = sheet.row_values(row)
    while (date[2] != u'低'):
        date = sheet.row_values(row + 1)
        row = row + 1
    document = Document('report.docx')
    document.add_page_break()
    table = document.add_table(rows=row - 1, cols=6)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = u'序号'
    hdr_cells[1].text = u'危险程度'
    hdr_cells[2].text = u'漏洞名称'
    hdr_cells[3].text = u'影响IP'
    hdr_cells[4].text = u'出现次数'
    hdr_cells[5].text = u'CVEID'

    for i in xrange(1, row):

        date = sheet.row_values(i + 1)
        if(date[2] == u'低'):
            break

        row_cells = table.rows[i].cells
        row_cells[0].text = str(i)
        row_cells[1].text = date[2]
        row_cells[2].text = date[3]
        row_cells[3].text = str(date[4])
        row_cells[4].text = str(date[5])
        row_cells[5].text = str(date[6])

    document.save(u'report.docx')


def doc_table(table_list):
    document = Document('report.docx')
    table = document.add_table(rows=7, cols=5)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = u'扫描对象'
    hdr_cells[1].text = u'扫描数量'
    hdr_cells[2].text = u'漏洞总数'
    hdr_cells[3].text = u'高危漏洞'
    hdr_cells[4].text = u'中危漏洞'
    hd_cells = table.columns[0].cells
    hd_cells[1].text = u'服务器'
    hd_cells[2].text = u'网络设备'
    hd_cells[3].text = u'终端设备'
    hd_cells[4].text = u'服务器'
    hd_cells[5].text = u'网络设备'
    hd_cells[6].text = u'终端设备'

    for i in xrange(6):
        row_cells = table.rows[i + 1].cells
        row_cells[1].text = str(table_list[4 * i])
        row_cells[2].text = str(table_list[4 * i + 1])
        row_cells[3].text = str(table_list[4 * i + 2])
        row_cells[4].text = str(table_list[4 * i + 3])
    document.save(u'report.docx')


def file_name(file_dir):
    L = []
    file_num = 0
    for root, dirs, files in os.walk(file_dir):
        for file in files:
            if file.endswith(".zip"):
                file_num = file_num + 1
                file_name = os.path.join(root, file)
                L.append(file_name)
    if(file_num != 6):
        print('zip file error：zip file num ' + str(file_num) + ',check zip file！')
        os._exit(0)
    return L


def un_zip(file_name, name):
    """unzip zip file"""

  #  file_name=raw_input(file_name1)
    try:
        zip_file = zipfile.ZipFile(file_name)

        # for names in zip_file.namelist():
        zip_file.extract('index.xls', name)
        zip_file.close()
    except KeyError:
        print('no such file index.xls ')
        os._exit(0)


def read_excel_4(list):
    temp1 = (u'终端总数:__' + str(list[8]) + u'__扫描终端个数:__' + str(list[8]) + u'__' + u'\n' +
             u'服务器总数:__' + str(list[0]) + u'__扫描服务器个数:__' + str(list[0]) + u'__' + u'\n' +
             u'数据库总数:__' + str(list[12]) + u'__扫描数据库个数:__' + str(list[12]) + u'__' + u'\n' +
             u'网络设备总数:__' + str(list[4]) + u'__扫描网络设备个数:__' + str(list[4]) + u'__' + u'\n' +
             u'发现漏洞总数:__' + str(list[1] + list[9]) + u'__' + u'\n' +
             u'终端漏洞个数：高__' + str(list[10]) + u'__中__' + str(list[11]) + u'__低__' + u'0' + u'__\n' +
             u'服务器漏洞个数：高__' + str(list[2]) + u'__中_' + str(list[3]) + u'__低__' + u'0' + u'__\n' +
             u'数据库漏洞个数：高__' + str(list[13]) + u'__中__' + str(list[14]) + u'__低__' + u'0' + u'__\n' +
             u'网络设备漏洞个数：高__' + str(list[5]) + u'__中__' + str(list[6]) + u'__低__' + u'0' + u'__\n')
    return temp1


def read_excel_3(excel_name):
    output_word = []
    book = xlrd.open_workbook(excel_name)
    sheet_name = book.sheet_names()[0]  # 获得指定索引的sheet名字
    sheet = book.sheet_by_name(sheet_name)
    station_num = int((sheet.cell_value(2, 1)))
    sheet_name = book.sheet_names()[1]
    sheet = book.sheet_by_name(sheet_name)
    hrisk_sum = int(sum(sheet.col_values(5)[2:30]))
    mrisk_sum = int(sum(sheet.col_values(6)[2:30]))
    risk_sum = hrisk_sum + mrisk_sum
    output_word.append(station_num)
    output_word.append(risk_sum)
    output_word.append(hrisk_sum)
    output_word.append(mrisk_sum)
    return output_word


def read_excel_22(excel_name):
    output_word = []
    safe_report = []
    book = xlrd.open_workbook(excel_name)
    sheet_name = book.sheet_names()[0]  # 获得指定索引的sheet名字
    sheet = book.sheet_by_name(sheet_name)
    station_num = int(sheet.cell_value(2, 1))
    sheet_name = book.sheet_names()[1]
    sheet = book.sheet_by_name(sheet_name)
    hrisk_col = int(sum(sheet.col_values(5)[2:50]))
    mrisk_col = int(sum(sheet.col_values(6)[2:50]))
    safe_report.append(u'本月共扫描外网桌面终端' + str(station_num) + u'台，')
    safe_report.append(u'发现' + str(hrisk_col) + u'个高危漏洞、' +
                       str(mrisk_col) + u'个中危漏洞，已整改高危漏洞3个，中危漏洞10个，其余已列入4月整改计划。')
    safe_report.append(u'发现' + str(mrisk_col) + u'个中危漏洞。具体扫描结果见附件。')
    safe_report.append(u'发现' + str(hrisk_col) + u'个高危漏洞。具体扫描结果见附件。')
    safe_report.append(u'无高、中危漏洞。具体扫描结果见附件。')
    if hrisk_col == 0 and mrisk_col == 0:
        output_word.append(safe_report[0])
        output_word.append(safe_report[4])
    if hrisk_col == 1 and mrisk_col == 1:
        output_word.append(safe_report[0])
        output_word.append(safe_report[1])
    if hrisk_col == 1 and mrisk_col == 0:
        output_word.append(safe_report[0])
        output_word.append(safe_report[3])
    if hrisk_col == 0 and mrisk_col == 1:
        output_word.append(safe_report[0])
        output_word.append(safe_report[2])

    return output_word


def read_excel_21(excel_name):

    safe_report = []
    book = xlrd.open_workbook(excel_name)
    sheet_name = book.sheet_names()[0]  # 获得指定索引的sheet名字
    sheet = book.sheet_by_name(sheet_name)
    station_num = int(sheet.cell_value(2, 1))
    sheet_name = book.sheet_names()[1]
    sheet = book.sheet_by_name(sheet_name)
    hrisk_col = int(sum(sheet.col_values(5)[2:50]))
    mrisk_col = int(sum(sheet.col_values(6)[2:50]))
    # del hrisk_col[0:1]
    # del mrisk_col[0:1]
    safe_report.append(u'本月共扫描内网终端' + str(station_num))
    safe_report.append(u'发现' + str(hrisk_col) + u'个高危、' +
                       str(mrisk_col) + u'个中危漏洞，已整改高危漏洞个，中危漏洞个，其余已列入本月整改计划。')
    return safe_report


def read_excel_1(excel_name):
    book = xlrd.open_workbook(excel_name)
    sheet_name = book.sheet_names()[1]  # 获得指定索引的sheet名字
    sheet = book.sheet_by_name(sheet_name)
    col_data = sheet.col_values(1)
    risk_data = []
    safe_report = []
    output_word = []
    try:
        row_portal_server = col_data.index('10.232.80.20')
        portal_server_hrisk = int(sheet.cell_value(row_portal_server, 5))
        portal_server_mrisk = int(sheet.cell_value(row_portal_server, 6))
        risk_data.append(portal_server_hrisk)
        risk_data.append(portal_server_mrisk)
    except ValueError:
        portal_server_hrisk = 0
        portal_server_mrisk = 0
        risk_data.append(portal_server_hrisk)
        risk_data.append(portal_server_mrisk)
    try:
        row_backup_server = col_data.index('10.232.80.1')
        backup_server_hrisk = int(sheet.cell_value(row_backup_server, 5))
        backup_server_mrisk = int(sheet.cell_value(row_backup_server, 6))
        risk_data.append(backup_server_hrisk)
        risk_data.append(backup_server_mrisk)

    except ValueError:
        backup_server_hrisk = 0
        backup_server_mrisk = 0
        risk_data.append(backup_server_hrisk)
        risk_data.append(backup_server_mrisk)

    try:
        row_dns_server = col_data.index('10.232.82.65')
        dns_server_hrisk = int(sheet.cell_value(row_dns_server, 5))
        dns_server_mrisk = int(sheet.cell_value(row_dns_server, 6))
        risk_data.append(dns_server_hrisk)
        risk_data.append(dns_server_mrisk)
    except ValueError:
        dns_server_hrisk = 0
        dns_server_mrisk = 0
        risk_data.append(dns_server_hrisk)
        risk_data.append(dns_server_mrisk)
    try:
        row_desk_server = col_data.index('10.232.80.54')
        desk_server_hrisk = int(sheet.cell_value(row_desk_server, 5))
        desk_server_mrisk = int(sheet.cell_value(row_desk_server, 6))
        risk_data.append(desk_server_hrisk)
        risk_data.append(desk_server_mrisk)
    except ValueError:
        desk_server_hrisk = 0
        desk_server_mrisk = 0
        risk_data.append(desk_server_hrisk)
        risk_data.append(desk_server_mrisk)
    try:
        row_deskbackup_server = col_data.index('10.232.80.53')
        deskbackup_server_hrisk = int(
            sheet.cell_value(row_deskbackup_server, 5))
        deskbackup_server_mrisk = int(
            sheet.cell_value(row_deskbackup_server, 6))
        risk_data.append(deskbackup_server_mrisk)
        risk_data.append(deskbackup_server_mrisk)
    except ValueError:
        deskbackup_server_hrisk = 0
        deskbackup_server_mrisk = 0
        risk_data.append(deskbackup_server_mrisk)
        risk_data.append(deskbackup_server_mrisk)
    if ((risk_data[0] or risk_data[1] or risk_data[2] or risk_data[3] or risk_data[4] or risk_data[5] or risk_data[6] or risk_data[7] or risk_data[8] or risk_data[9]) == 0):
        print ('risk num is zero!!!!')
        os._exit(0)
    safe_report.append(u'本月使用绿盟扫描，')
    safe_report.append(u'门户服务器')
    safe_report.append(str(risk_data[0]) + u'个高危漏洞（已备案），')
    safe_report.append(str(risk_data[1]) + u'个中危漏洞（已备案），')
    safe_report.append(u'DNS服务器')
    safe_report.append(str(risk_data[2]) + u'个高危漏洞（已备案），')
    safe_report.append(str(risk_data[3]) + u'个中危漏洞（已备案），')
    safe_report.append(u'备份服务器')
    safe_report.append(str(risk_data[4]) + u'个高危漏洞（已备案），')
    safe_report.append(str(risk_data[5]) + u'个中危漏洞（已备案），')
    safe_report.append(u'桌管服务器')
    safe_report.append(str(risk_data[6]) + u'个高危漏洞（已备案），')
    safe_report.append(str(risk_data[7]) + u'个中危漏洞（已备案），')
    safe_report.append(u'桌管备份服务器')
    safe_report.append(str(risk_data[8]) + u'个高危漏洞（已备案），')
    safe_report.append(str(risk_data[9]) + u'个中危漏洞（已备案），')
    safe_report.append(u'外网主机、服务器、终端均未发现中、高危漏洞。自建系统未发现高、中危漏洞。具体扫描结果见附件。')
    output_word.append(safe_report[0])
    for i in range(5):
        if risk_data[2 * i] != 0 or risk_data[2 * i + 1] != 0:
            output_word.append(safe_report[3 * i + 1])
        if risk_data[2 * i] != 0:
            output_word.append(safe_report[3 * i + 2])
        if risk_data[2 * i + 1] != 0:
            output_word.append(safe_report[3 * i + 3])
    output_word.append(safe_report[16])
    return output_word
def rm_dir():
    for num in range(6):
        shutil.rmtree(str(num))
    


if __name__ == "__main__":
    all_zip_file = file_name('./')
    root = "./"
    target_file = []
    document = Document(u'G:\code\Na_tool')
    document.save(u'report.docx')
    for num in range(6):
        un_zip(all_zip_file[num], str(num))
        name_temp = (root + str(num) + '\index.xls')
        target_file.append(name_temp)

# 安全月报第一段
    print ('safe report part 1 making ......')
    start_report1 = time.clock()
    safe_report_1 = read_excel_1(target_file[2])
    try:
        doc(safe_report_1)
    except IndexError:
        print('safe report part 1 data error')
    end_report1 = time.clock()
    print ( str(int((end_report1 - start_report1) * 1000)) + u'ms')
# 安全月报第二段
    print ('safe report part 2 making ......')
    start_report2 = time.clock()
    safe_report_21 = read_excel_21(target_file[1])
    safe_report_22 = read_excel_22(target_file[3])
    safe_report_2 = safe_report_21 + safe_report_22
    try:
        doc(safe_report_2)
    except IndexError:
        print('safe report part 2 data error')
    end_report2 = time.clock()
    print (str(int((end_report2 - start_report2) * 1000)) + u'ms')

# 安全月报第三段
    print ('safe report part 3 making ......')
    start_report3 = time.clock()
    safe_report_3 = read_excel_3(target_file[2])
    safe_report_3 = safe_report_3 + read_excel_3(target_file[0])
    safe_report_3 = safe_report_3 + read_excel_3(target_file[1])
    safe_report_3 = safe_report_3 + read_excel_3(target_file[5])
    safe_report_3 = safe_report_3 + read_excel_3(target_file[3])
    safe_report_3 = safe_report_3 + read_excel_3(target_file[4])
    try:
        doc_table(safe_report_3)
    except IndexError:
        print('safe report part 3 data error')
    end_report3 = time.clock()
    print ( str(int((end_report3 - start_report3) * 1000)) + u'ms')
    print("safe report finished")

# 督查月报
    print ('check report making......')
    start_report4 = time.clock()
    safe_report_4 = read_excel_4(safe_report_3)
    try:
        doc(safe_report_4)
    except IndexError:
        print('check report  data error')
    end_report4 = time.clock()
    print ( str(int((end_report4 - start_report4) * 1000)) + u'ms')
    print ('check report  finished')
# 终端统计表
    print ('station data making......')
    start_report5 = time.clock()
    try:
        doc_table_risk_data1(target_file[1])
    except IndexError:
        print('station date error')
    end_report5 = time.clock()
    print ( str(int((end_report5 - start_report5) * 1000)) + u'ms')
    print ('station data finished')
# 漏洞明细表
    print ('risk data making......')
    start_report6 = time.clock()
    try:
        doc_table_risk_data2(target_file[1])
    except IndexError:
        print('risk data error ')
    end_report6 = time.clock()
    print ( str(int((end_report6 - start_report6) * 1000)) + u'ms')
    print ('risk data finished')
    rm_dir()
    